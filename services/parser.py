from pypdf import PdfReader
from docx import Document
import os


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using pypdf"""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
    return text.strip()


def extract_text(file_path):
    """Extract text from a file based on its extension"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def allowed_file(filename, allowed_extensions):
    """Check if the file has an allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions
